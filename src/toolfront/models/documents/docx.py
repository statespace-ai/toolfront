import logging

from toolfront.models.documents.base import Document, DocumentError

logger = logging.getLogger("toolfront")


class DOCXDocument(Document):
    """DOCX document reader."""

    async def read(self, pagination: int | float = 0) -> str:
        """Read DOCX document content.

        Args:
            pagination: Ignored for DOCX documents.

        Returns:
            Document content as string.
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(self.path)

            # Read all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text
        except ImportError:
            error_msg = "DOCX support requires python-docx library"
            logger.error(f"Missing dependency for DOCX reading: {error_msg}")
            raise DocumentError(error_msg)
        except Exception as e:
            logger.error(f"Error reading DOCX file {self.url}: {e}")
            raise DocumentError(f"Error reading DOCX file: {str(e)}")
